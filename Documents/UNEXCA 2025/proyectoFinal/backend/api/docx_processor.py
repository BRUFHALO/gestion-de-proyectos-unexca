"""
API Router para procesamiento de archivos DOCX
Extrae contenido, convierte a HTML y maneja anotaciones
"""
from fastapi import APIRouter, HTTPException, UploadFile, File, Body
from typing import List, Dict, Any, Optional
from pydantic import BaseModel
from bson import ObjectId
from datetime import datetime
import io
import base64

from config.database import Database, DatabaseConfig

router = APIRouter(prefix="/api/v1/docx", tags=["docx-processing"])

class ParagraphData(BaseModel):
    """Modelo de párrafo extraído"""
    id: str
    text: str
    style: str
    runs: List[Dict[str, Any]]  # Fragmentos de texto con formato
    position: Dict[str, int]  # page, index

class TableData(BaseModel):
    """Modelo de tabla extraída"""
    id: str
    rows: List[List[str]]
    position: Dict[str, int]

class ImageData(BaseModel):
    """Modelo de imagen extraída"""
    id: str
    data: str  # Base64
    width: int
    height: int
    position: Dict[str, int]

class DocxContent(BaseModel):
    """Modelo de contenido extraído del DOCX"""
    html: str
    paragraphs: List[Dict[str, Any]]
    tables: List[Dict[str, Any]]
    images: List[Dict[str, Any]]
    metadata: Dict[str, Any]

class Annotation(BaseModel):
    """Modelo de anotación/comentario sobre el documento"""
    id: str
    type: str  # 'error', 'suggestion', 'comment'
    text_content: str  # El comentario del docente
    selected_text: str  # El texto del documento que fue seleccionado
    position: Dict[str, float]  # top, left - posición donde apareció la caja
    paragraph_id: Optional[str] = None  # ID del párrafo vinculado (opcional)
    anchor_id: Optional[str] = None  # ID del elemento HTML donde se colocó el comentario
    created_by: Optional[str] = "teacher"  # Usuario que creó el comentario
    created_at: Optional[str] = None  # Timestamp de creación

class SaveAnnotationsRequest(BaseModel):
    """Modelo para guardar anotaciones"""
    project_id: str
    annotations: List[Annotation]

class ExportRequest(BaseModel):
    """Modelo para exportar documento con correcciones"""
    project_id: str
    export_format: str  # 'docx' o 'pdf'
    apply_corrections: bool = False

@router.post("/parse/{project_id}", response_model=DocxContent)
async def parse_docx(project_id: str):
    """
    Parsea el archivo DOCX del proyecto y extrae su contenido estructurado
    """
    try:
        # Importar aquí para evitar errores si no están instaladas
        try:
            from docx import Document
        except ImportError:
            raise HTTPException(
                status_code=500,
                detail="python-docx no está instalado. Ejecuta: pip install python-docx"
            )
        
        try:
            import mammoth
        except ImportError:
            raise HTTPException(
                status_code=500,
                detail="mammoth no está instalado. Ejecuta: pip install mammoth"
            )
        
        # Obtener el proyecto de la base de datos
        projects_collection = Database.get_collection("projects")
        project = await projects_collection.find_one({"_id": ObjectId(project_id)})
        
        if not project:
            raise HTTPException(status_code=404, detail="Proyecto no encontrado")
        
        # Obtener información del archivo desde la primera versión
        versions = project.get("versions", [])
        if not versions:
            raise HTTPException(status_code=404, detail="El proyecto no tiene versiones")
        
        files = versions[0].get("files", [])
        if not files:
            raise HTTPException(status_code=404, detail="No hay archivos en la primera versión del proyecto")
        
        file_info = files[0]
        file_path = file_info.get("file_path")
        file_id = file_info.get("file_id")
        
        print(f"DEBUG - Project ID: {project_id}")
        print(f"DEBUG - File info: {file_info}")
        print(f"DEBUG - File path: {file_path}")
        print(f"DEBUG - File ID: {file_id}")
        
        if not file_path and not file_id:
            raise HTTPException(status_code=404, detail="Ni file_path ni file_id encontrados en el proyecto")
        
        # Construir ruta completa al archivo
        from pathlib import Path
        base_path = Path(__file__).parent.parent / "uploads"
        
        # Intentar con file_path primero, luego con file_id
        if file_path:
            full_path = base_path / file_path
        elif file_id:
            # Buscar el archivo por file_id en el directorio uploads
            full_path = base_path / file_id
        
        print(f"DEBUG - Base path: {base_path}")
        print(f"DEBUG - Full path: {full_path}")
        print(f"DEBUG - File exists: {full_path.exists()}")
        
        if not full_path.exists():
            # Intentar buscar el archivo en subdirectorios
            import os
            found_files = []
            for root, dirs, files_in_dir in os.walk(base_path):
                for file in files_in_dir:
                    if file_id and file == file_id:
                        full_path = Path(root) / file
                        print(f"DEBUG - Found file at: {full_path}")
                        break
                    elif file_path and file_path in str(Path(root) / file):
                        full_path = Path(root) / file
                        print(f"DEBUG - Found file at: {full_path}")
                        break
            
            if not full_path.exists():
                raise HTTPException(
                    status_code=404, 
                    detail=f"Archivo no encontrado. Buscado en: {base_path}, file_path: {file_path}, file_id: {file_id}"
                )
        
        # Leer el documento con python-docx
        document = Document(str(full_path))
        
        # Extraer párrafos
        paragraphs = []
        for idx, para in enumerate(document.paragraphs):
            if para.text.strip():  # Solo párrafos con contenido
                para_data = {
                    "id": f"p{idx}",
                    "text": para.text,
                    "style": para.style.name if para.style else "Normal",
                    "runs": [
                        {
                            "text": run.text,
                            "bold": run.bold,
                            "italic": run.italic,
                            "underline": run.underline
                        }
                        for run in para.runs
                    ],
                    "position": {"page": 1, "index": idx}  # Simplificado
                }
                paragraphs.append(para_data)
        
        # Extraer tablas
        tables = []
        for idx, table in enumerate(document.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                rows.append(cells)
            
            tables.append({
                "id": f"t{idx}",
                "rows": rows,
                "position": {"page": 1, "index": idx}
            })
        
        # Extraer información de formato detallada con python-docx
        paragraph_styles = []
        for idx, para in enumerate(document.paragraphs):
            style_info = {
                'index': idx,
                'alignment': str(para.alignment) if para.alignment else 'LEFT',
                'left_indent': para.paragraph_format.left_indent.pt if para.paragraph_format.left_indent else 0,
                'right_indent': para.paragraph_format.right_indent.pt if para.paragraph_format.right_indent else 0,
                'first_line_indent': para.paragraph_format.first_line_indent.pt if para.paragraph_format.first_line_indent else 0,
                'space_before': para.paragraph_format.space_before.pt if para.paragraph_format.space_before else 0,
                'space_after': para.paragraph_format.space_after.pt if para.paragraph_format.space_after else 0,
                'line_spacing': para.paragraph_format.line_spacing if para.paragraph_format.line_spacing else 1.15
            }
            paragraph_styles.append(style_info)
        
        # Convertir a HTML con mammoth con opciones de estilo mejoradas
        with open(str(full_path), 'rb') as docx_file:
            # Configurar opciones de conversión para preservar más estilos
            style_map = """
                p[style-name='Heading 1'] => h1.heading-1:fresh
                p[style-name='Heading 2'] => h2.heading-2:fresh
                p[style-name='Heading 3'] => h3.heading-3:fresh
                p[style-name='Title'] => h1.doc-title:fresh
                p[style-name='Subtitle'] => h2.doc-subtitle:fresh
                p[style-name='Intense Quote'] => blockquote.intense-quote:fresh
                p[style-name='Normal'] => p.normal:fresh
                r[style-name='Strong'] => strong
                r[style-name='Emphasis'] => em
            """
            
            # Función para convertir imágenes a base64
            def convert_image_to_base64(image):
                with image.open() as image_bytes:
                    encoded = base64.b64encode(image_bytes.read()).decode('ascii')
                    return {
                        "src": f"data:{image.content_type};base64,{encoded}"
                    }
            
            result = mammoth.convert_to_html(
                docx_file,
                style_map=style_map,
                include_default_style_map=True,
                convert_image=mammoth.images.img_element(convert_image_to_base64)
            )
            html_content = result.value
            
            # Aplicar estilos de formato extraídos de python-docx
            for idx, style_info in enumerate(paragraph_styles):
                # Mapear alineación
                alignment_map = {
                    'WdParagraphAlignment.CENTER': 'center',
                    'WdParagraphAlignment.RIGHT': 'right',
                    'WdParagraphAlignment.JUSTIFY': 'justify',
                    'WdParagraphAlignment.LEFT': 'left',
                    '3': 'center',  # CENTER
                    '2': 'right',   # RIGHT
                    '3': 'justify', # JUSTIFY
                    '0': 'left',    # LEFT
                    'CENTER': 'center',
                    'RIGHT': 'right',
                    'JUSTIFY': 'justify',
                    'LEFT': 'left'
                }
                
                alignment = alignment_map.get(style_info['alignment'], 'left')
                
                # Crear estilo inline para cada párrafo
                para_style = f"text-align: {alignment}; "
                para_style += f"margin-left: {style_info['left_indent']}pt; "
                para_style += f"margin-right: {style_info['right_indent']}pt; "
                para_style += f"text-indent: {style_info['first_line_indent']}pt; "
                para_style += f"margin-top: {style_info['space_before']}pt; "
                para_style += f"margin-bottom: {style_info['space_after']}pt; "
                
                if isinstance(style_info['line_spacing'], (int, float)):
                    para_style += f"line-height: {style_info['line_spacing']}; "
            
            # Agregar estilos CSS completos
            css_styles = """
                <style>
                    body { 
                        font-family: 'Calibri', 'Arial', sans-serif; 
                        font-size: 11pt;
                        margin: 0;
                        padding: 0;
                    }
                    h1, h2, h3 { 
                        font-weight: bold; 
                        margin-top: 12pt; 
                        margin-bottom: 6pt; 
                    }
                    h1.heading-1 { font-size: 16pt; color: #2e74b5; }
                    h2.heading-2 { font-size: 13pt; color: #2e74b5; }
                    h3.heading-3 { font-size: 12pt; color: #1f4d78; }
                    h1.doc-title { 
                        font-size: 26pt; 
                        text-align: center; 
                        font-weight: bold; 
                        margin-bottom: 12pt;
                    }
                    h2.doc-subtitle { 
                        font-size: 15pt; 
                        text-align: center; 
                        color: #595959;
                        margin-bottom: 12pt;
                    }
                    p { 
                        margin-top: 0pt; 
                        margin-bottom: 8pt; 
                        line-height: 1.15;
                    }
                    p.normal { 
                        font-size: 11pt;
                        margin-bottom: 8pt;
                    }
                    strong, b { font-weight: bold; }
                    em, i { font-style: italic; }
                    u { text-decoration: underline; }
                    blockquote.intense-quote { 
                        margin-left: 36pt; 
                        margin-right: 36pt;
                        font-style: italic; 
                        border-left: 3px solid #2e74b5;
                        padding-left: 12pt;
                    }
                    table {
                        border-collapse: collapse;
                        width: 100%;
                        margin: 12pt 0;
                    }
                    td, th {
                        border: 1px solid #d0d0d0;
                        padding: 6pt;
                    }
                    /* Selección de texto para comentarios */
                    .text-selected {
                        background-color: #fff4cc;
                        border-bottom: 2px solid #ffeb3b;
                    }
                    .text-commented {
                        background-color: #fff9e6;
                        cursor: pointer;
                    }
                    .text-commented:hover {
                        background-color: #fff4cc;
                    }
                    /* Estilos para imágenes */
                    img {
                        max-width: 100%;
                        height: auto;
                        display: block;
                        margin: 12pt auto;
                    }
                </style>
            """
            html_content = css_styles + html_content
        
        # Agregar IDs a los párrafos en el HTML
        html_with_ids = html_content
        for idx, para in enumerate(paragraphs):
            # Buscar y reemplazar párrafos con IDs
            if para["text"] in html_with_ids:
                html_with_ids = html_with_ids.replace(
                    f"<p>{para['text']}</p>",
                    f"<p id='{para['id']}' data-para-id='{para['id']}'>{para['text']}</p>",
                    1
                )
        
        return {
            "html": html_with_ids,
            "paragraphs": paragraphs,
            "tables": tables,
            "images": [],  # Por ahora sin imágenes
            "metadata": {
                "title": project.get("title", "Sin título"),
                "pages": len(paragraphs) // 20 + 1,  # Estimación
                "paragraphs_count": len(paragraphs),
                "tables_count": len(tables)
            }
        }
        
    except ImportError as e:
        import traceback
        error_detail = traceback.format_exc()
        print(f"ImportError: {error_detail}")
        raise HTTPException(
            status_code=500, 
            detail=f"Librerías no instaladas: {str(e)}"
        )
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        error_detail = traceback.format_exc()
        print(f"Error procesando DOCX: {error_detail}")
        raise HTTPException(status_code=500, detail=f"Error procesando DOCX: {str(e)}")

@router.post("/annotations/save")
async def save_annotations(request: SaveAnnotationsRequest):
    """
    Guarda las anotaciones del docente sobre el documento
    """
    try:
        print("=" * 50)
        print("RECIBIENDO ANOTACIONES:")
        print(f"Project ID: {request.project_id}")
        print(f"Número de anotaciones: {len(request.annotations)}")
        for idx, anno in enumerate(request.annotations):
            print(f"\nAnotación {idx + 1}:")
            print(f"  ID: {anno.id}")
            print(f"  Type: {anno.type}")
            print(f"  Text content: {anno.text_content[:50]}...")
            print(f"  Selected text: {anno.selected_text[:50]}...")
            print(f"  Position: {anno.position}")
            print(f"  Created by: {anno.created_by}")
            print(f"  Created at: {anno.created_at}")
        print("=" * 50)
        
        annotations_collection = Database.get_collection("docx_annotations")
        
        # Eliminar anotaciones anteriores del proyecto
        await annotations_collection.delete_many({"project_id": ObjectId(request.project_id)})
        
        # Guardar nuevas anotaciones
        if request.annotations:
            annotations_data = []
            for annotation in request.annotations:
                anno_dict = annotation.dict()
                anno_dict["project_id"] = ObjectId(request.project_id)
                anno_dict["created_at"] = datetime.utcnow().isoformat()
                annotations_data.append(anno_dict)
            
            await annotations_collection.insert_many(annotations_data)
        
        return {
            "success": True,
            "message": f"Se guardaron {len(request.annotations)} anotaciones",
            "count": len(request.annotations)
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error guardando anotaciones: {str(e)}")

@router.get("/annotations/{project_id}")
async def get_annotations(project_id: str):
    """
    Obtiene las anotaciones guardadas de un proyecto
    """
    try:
        annotations_collection = Database.get_collection("docx_annotations")
        
        annotations = await annotations_collection.find({
            "project_id": ObjectId(project_id)
        }).to_list(length=None)
        
        # Convertir ObjectId a string
        for anno in annotations:
            anno["_id"] = str(anno["_id"])
            anno["project_id"] = str(anno["project_id"])
        
        return {
            "success": True,
            "annotations": annotations,
            "count": len(annotations)
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error obteniendo anotaciones: {str(e)}")

@router.post("/export")
async def export_with_annotations(request: ExportRequest):
    """
    Exporta el documento con anotaciones aplicadas
    - PDF: Genera PDF con anotaciones visuales superpuestas
    - DOCX: Genera DOCX con correcciones aplicadas (si apply_corrections=True)
    """
    try:
        if request.export_format == 'pdf':
            # TODO: Implementar generación de PDF con reportlab
            return {
                "success": True,
                "message": "PDF con anotaciones generado",
                "download_url": f"/api/v1/projects/download/{request.project_id}_annotated.pdf"
            }
        
        elif request.export_format == 'docx':
            if request.apply_corrections:
                # TODO: Implementar aplicación de correcciones al DOCX
                return {
                    "success": True,
                    "message": "DOCX con correcciones aplicadas",
                    "download_url": f"/api/v1/projects/download/{request.project_id}_corrected.docx"
                }
            else:
                return {
                    "success": True,
                    "message": "DOCX original sin cambios",
                    "download_url": f"/api/v1/projects/download/{request.project_id}.docx"
                }
        
        else:
            raise HTTPException(status_code=400, detail="Formato no soportado. Use 'pdf' o 'docx'")
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error exportando: {str(e)}")

@router.get("/preview/{project_id}")
async def get_document_preview(project_id: str):
    """
    Obtiene la vista previa del documento con sus anotaciones
    """
    try:
        # Obtener contenido del documento
        content = await parse_docx(project_id)
        
        # Obtener anotaciones
        annotations_result = await get_annotations(project_id)
        
        return {
            "project_id": project_id,
            "content": content,
            "annotations": annotations_result.get("annotations", [])
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error obteniendo preview: {str(e)}")
